import customtkinter as ctk
import pandas as pd
from tkinter import messagebox as mx
from quick_analysis.scraping import descarga_info, scrapeo
from quick_analysis.data_transformation import limpiar_df, analisis_df
from config import Config
from quick_analysis.transformation import progress, buscar_archivos
from docx import Document
import os
import time




def option_selected_year (option):
    print(f"Has seleccionado el año: {option}")
    global año
    año = int(option)

def option_selected_div (option):
    print(f"Has seleccionado la division: {option}")
    global div
    div = int(option)

def condicional_interruptor(scraping_window):
    global frame_button, button_continue, label_año, combo_year, label_division, combo_div, frame_button_divisiones, button_continue_divisiones

    if switch_var.get() == 1:
        switch.configure(text="Scrapeo Parcial")
        if frame_button is not None and button_continue is not None:
            frame_button.destroy()
            button_continue.destroy()
            frame_button = None
            button_continue = None
        
        
        if label_año is None and combo_year is None and label_division is None and combo_div is None and frame_button_divisiones is None and button_continue_divisiones is None:
            # año
            label_año = ctk.CTkLabel(scraping_window, text="Seleccione un año", font=("Helvetica", 15))
            label_año.pack(pady=5, padx=5)

            # Crear un ComboBox con opciones
            combo_year = ctk.CTkComboBox(scraping_window, 
                                        values=["1","2","3","4","5"],
                                        command=option_selected_year)  # Se ejecuta al seleccionar
            combo_year.pack(pady=5)

            # Establecer un valor por defecto
            combo_year.set("Seleccionar año")

            # Seleccionar division
            label_division = ctk.CTkLabel(scraping_window, text="Seleccionar division", font=("Helvetica", 15))
            label_division.pack(pady=5, padx=5)

            # Crear un ComboBox con opciones
            combo_div = ctk.CTkComboBox(scraping_window, 
                                        values=["1","2","3","4","5"],
                                        command=option_selected_div)  # Se ejecuta al seleccionar
            combo_div.pack(pady=5)

            # Establecer un valor por defecto
            combo_div.set("Selecciona año")
            frame_button_divisiones = ctk.CTkFrame(scraping_window)
            frame_button_divisiones.pack(pady=25, padx=5)
            button_continue_divisiones = ctk.CTkButton(frame_button_divisiones, text="Continue", command=lambda: main(year=año,div=div,switch_var= switch_var.get()))
            button_continue_divisiones.pack(pady=2, padx=2)
    else:
        switch.configure(text="Scrapeo completo")
        if label_año is not None and combo_year is not None and label_division is not None and combo_div is not None:
            label_año.destroy()
            combo_year.destroy()
            label_division.destroy()
            combo_div.destroy()
            frame_button_divisiones.destroy()
            button_continue_divisiones.destroy()
            label_año = None
            combo_year = None
            label_division = None
            combo_div = None
            frame_button_divisiones=None
            button_continue_divisiones = None

        if frame_button is None and button_continue is None:
            frame_button = ctk.CTkFrame(scraping_window)
            frame_button.pack(pady=5, padx=5)
            button_continue = ctk.CTkButton(frame_button, text="Continue", command=lambda: main(switch_var=switch_var.get()))
            button_continue.pack(pady=2, padx=2)

 

def fuction_scraping_window():
    ## Generamos la ventana que se ejecutara cuando apretemos el boton de analisis rapido
    scraping_window = ctk.CTk()
    scraping_window.title("Ventana Scrapeo")
    scraping_window.geometry("850x450")


    # Condicionamos a que si el scrapeo va a ser completo o no con un switch
    # Declaramos todas las variables antes de crearlas
    global frame_button, button_continue, label_año, combo_year, label_division, combo_div, frame_button_divisiones, button_continue_divisiones

    label_año = None
    combo_year = None
    label_division = None
    combo_div = None
    frame_button_divisiones = None
    button_continue_divisiones = None

    frame_button = None
    button_continue = None
    # Variable asociada al interruptor
    global switch_var
    switch_var = ctk.IntVar()

    # Crear interruptor
    global switch
    switch = ctk.CTkSwitch(scraping_window, text="Seleccione una modalidad", variable=switch_var, command=lambda:condicional_interruptor(scraping_window), onvalue=0, offvalue=1)
    switch.pack(pady=15, padx=2)
    

    scraping_window.mainloop()



def cuenta_regresiva(root, label_texto, label_numerico, i=5, callback=None):
    if i > 0:
        label_numerico.configure(text=f"\n\n{i}")
        root.after(1000, cuenta_regresiva, root, label_texto,label_numerico, i-1, callback)  # Llamar a la función cada segundo
    else:
        label_texto.destroy()
        label_numerico.configure(text="\n\n¡Iniciando automatización!\n¡Suelte los periféricos!", font=("Helvetica", 32))
        root.after(5000, lambda: cerrar_ventana_y_ejecutar(root, callback))  # Esperar 5 segundos antes de cerrar

def cerrar_ventana_y_ejecutar(root, callback):
    root.destroy()  # Cerrar la ventana
    if callback:
        callback()  # Ejecutar la automatización

def main(year=None, div=None, switch_var=None):
    # Creamos la ventana
    warning_window = ctk.CTk()
    warning_window.title("Comienzo automatización")
    warning_window.geometry("750x450")

    # Creamos el label con el mensaje de cuenta regresiva
    label_cuenta_regresiva = ctk.CTkLabel(warning_window, text="\n\nLa automatización comenzará en: ", font=("Helvetica", 32))
    label_cuenta_regresiva.pack(pady=5, padx=5)
    label_numerico = ctk.CTkLabel(warning_window, text="\n\n5", font=("Helvetica", 48))
    label_numerico.pack(pady=5, padx=5)

    # Iniciar la cuenta regresiva con una función callback que se ejecuta al finalizar
    cuenta_regresiva(warning_window, label_cuenta_regresiva, label_numerico, 5, lambda: iniciar_proceso(year, div, switch_var))

    # Ejecutar la ventana
    warning_window.mainloop()

def iniciar_proceso(year, div, switch_var):
    """Esta función ejecuta todo el código después de la cuenta regresiva"""
    scrapeo(year, div, switch_var)

    try:
        archivos_descargados = buscar_archivos(Config.ruta_directorio_input)

        for archivo in archivos_descargados:
            # Luego de terminado el proceso de transformación de xls a xlsx extraemos el nombre resultado
            nombre_excel = progress(archivo)

            df = pd.read_excel(f"{Config.ruta_directorio_input}/{nombre_excel}")
            df_limpio = limpiar_df(df)
            resultado = analisis_df(df_limpio)

            # Guardar en Excel
            resultado.to_excel(f"{Config.ruta_salida}/xlsx/{nombre_excel}.xlsx", engine="openpyxl", index=False)

            try:
                nombre_word = nombre_excel.replace(".xlsx", ".docx")
                df_a_word = resultado[["Nombre y Apellido", "Cantidad materias"]]

                # Crear documento Word
                df_word = Document()
                df_word.add_heading(nombre_excel, level=1)

                # Crear tabla
                tabla = df_word.add_table(rows=1, cols=len(df_a_word.columns))
                encabezados = tabla.rows[0].cells
                for i, columna in enumerate(df_a_word.columns):
                    encabezados[i].text = columna

                for _, fila in df_a_word.iterrows():
                    row_cells = tabla.add_row().cells
                    for i, valor in enumerate(fila):
                        row_cells[i].text = str(valor)

                df_word.save(f"{Config.ruta_salida}/docx/{nombre_word}.docx")
                print(f"✅ Archivo Word '{nombre_word}' generado con éxito.")
            except Exception as e:
                print("❌ ERROR: al intentar exportar a Word", e)

        try:
            # Remover archivos de input
            archivos_xlsx_input = os.listdir(Config.ruta_directorio_input)
            for archivo in archivos_xlsx_input:
                os.remove(f"{Config.ruta_directorio_input}/{archivo}")
                print("🗑 Se removió con éxito el elemento de INPUTS")

        except Exception as e:
            print("❌ ERROR: Se produjo un error al intentar eliminar el elemento de INPUTS. Msj Error: ", e)

        mx.showinfo("Filtrado completo", "Se han exportado todos los archivos a la carpeta outputs")

    except Exception as e:
        print("❌ ERROR: Al exportar el excel", e)
